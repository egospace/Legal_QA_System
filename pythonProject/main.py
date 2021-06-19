import docx
import json
import re
from loguru import logger
import win32com.client as wc
import shutil
import os

logger.add("log/interface_log_{time}.log", rotation="500MB", encoding="utf-8", enqueue=True, compression="zip",
           retention="10 days")


def _trans(s):
    digit = {'一': 1, '二': 2, '三': 3, '四': 4, '五': 5, '六': 6, '七': 7, '八': 8, '九': 9}
    num = 0
    if s:
        idx_q, idx_b, idx_s = s.find('千'), s.find('百'), s.find('十')
        if idx_q != -1:
            num += digit[s[idx_q - 1:idx_q]] * 1000
        if idx_b != -1:
            num += digit[s[idx_b - 1:idx_b]] * 100
        if idx_s != -1:
            # 十前忽略一的处理
            num += digit.get(s[idx_s - 1:idx_s], 1) * 10
        if s[-1] in digit:
            num += digit[s[-1]]
    return num


# 事实x理由
def data_clean_ly():
    # lines = []
    # with open("法律种类.txt", 'r', encoding='utf-8') as fp:
    #     lines = fp.readlines()
    #     for i in range(len(lines)):
    #         lines[i] = re.findall(r'[《](.*?)[》]', lines[i])[0]
    #     fp.close()
    newfileNames = os.listdir(r"dataxf/查明/")
    l = 1
    for file in newfileNames:
        # 注意SaveAs会打开保存后的文件，有时可能看不到，但后台一定是打开的
        path = "dataxf/查明/" + \
               file.split(".docx")[0] + ".docx"
        res_path = "xiao.txt"
        print(file)
        fact = ""
        laws = []
        # print(path)
        # 获取文档的所有段落 path : 相对路径包含文档名称
        docx_temp = docx.Document(path)
        f = False
        for para in docx_temp.paragraphs:
            if re.findall("查明(.*)", str(para.text)):
                # key1 = re.findall("事实(.)理由", str(para.text))[0]
                fact = str(para.text).split("查明")[1]
                if len(fact) <= 20:
                    f = True
                    continue
                # print(str(len(fact)) + "<::>" + str(l) + "<:>" + fact)
            if f:
                fact = str(para.text)
                f = False
                # print(str(len(fact)) + "<::>" + str(l) + "<:>" + fact)
            if re.findall("(.*)判决如下", str(para.text)):
                t = re.findall("(.*)判决如下", str(para.text))[0]
                print(t)
                law_list = re.findall(r'《中华人民共和国消费者权益保护法》(.*?)[《]', t)
                temp = True
                if len(law_list) > 0:
                    temp = False
                if temp:
                    law_list = re.findall(r'《中华人民共和国消费者权益保护法》(.*?)规定', t)
                print(law_list)
                if len(law_list) == 0:
                    with open("dataraw/error.txt", 'a+', encoding='utf-8') as fp:
                        fp.write(file + '\n')
                        fp.close()
                else:
                    law_list = str(law_list[0]).split("、")
                    print(law_list)
                    for i in range(len(law_list)):
                        try:
                            laws.append(_trans(re.findall("第(.*)条", law_list[i])[0]))
                            # laws.append(_trans(re.findall("(.)十(.)条", law_list[i])[0]))
                        except:
                            continue
                    # if len(laws) == 0:
                    #     with open("keywords/诉称null.txt", 'a+', encoding='utf-8') as fp:
                    #         fp.write(str(file) + '\n')
                    #         fp.close()
                    #     logger.info("[" + str(l) + "]" + str(file))
                        # print(str(l) + "<::>" + str(laws))
                        # l += 1
                break
                # for i in range(len(law_list)):
                #     if law_list[i] in lines:
                #         laws.append(law_list[i])
        logger.info(str(l) + "<:>" + str(file))
        l += 1
        res = {'fact': fact, 'laws': laws}
        res2json = json.dumps(res, ensure_ascii=False)
        with open(res_path, 'a+', encoding='utf-8') as fp:
            fp.write(res2json + '\n')
            fp.close()


# 本院查明
def data_clean():
    lines = []
    with open("法律种类.txt", 'r', encoding='utf-8') as fp:
        lines = fp.readlines()
        for i in range(len(lines)):
            lines[i] = re.findall(r'[《](.*?)[》]', lines[i])[0]
        fp.close()
    newfileNames = os.listdir(r"D:\\Code\\pythoncode\\pythonProject\\dataraw\\本院查明\\根据x判决如下\\")
    l = 1
    for file in newfileNames:
        # 注意SaveAs会打开保存后的文件，有时可能看不到，但后台一定是打开的
        path = "D:\\Code\\pythoncode\\pythonProject\\dataraw\\本院查明\\根据x判决如下.txt\\" + file.split(".docx")[0] + ".docx"
        res_path = "res.txt"
        fact = ""
        laws = []
        # print(path)
        # 获取文档的所有段落 path : 相对路径包含文档名称
        docx_temp = docx.Document(path)
        f = False
        for para in docx_temp.paragraphs:
            if re.findall("本院查明(.*)", str(para.text)):
                # key1 = re.findall("事实(.)理由", str(para.text))[0]
                fact = str(para.text).split("本院查明")[1]
                if len(fact) <= 20:
                    f = True
                    continue
                # print(str(len(fact))+"<::>"+str(l)+"<:>"+fact)
            if f:
                fact = str(para.text)
                f = False
                # print(str(len(fact)) + "<::>" + str(l) + "<:>" + fact)
            if re.findall("根据(.*)判决如下", str(para.text)):
                t = re.findall("根据(.*)判决如下", str(para.text))[0]
                law_list = re.findall(r'[《](.*?)[》]', t)
                for i in range(len(law_list)):
                    if law_list[i] in lines:
                        laws.append(law_list[i])
        logger.info(str(l) + "<:>" + str(file))
        l += 1
        res = {'fact': fact, 'laws': laws}
        res2json = json.dumps(res, ensure_ascii=False)
        with open(res_path, 'a+', encoding='utf-8') as fp:
            fp.write(res2json + '\n')
            fp.close()


# 经审理查明
def data_clean_sl():
    lines = []
    with open("法律种类.txt", 'r', encoding='utf-8') as fp:
        lines = fp.readlines()
        for i in range(len(lines)):
            lines[i] = re.findall(r'[《](.*?)[》]', lines[i])[0]
        fp.close()
    # newfileNames = os.listdir(r"D:\\Code\\pythoncode\\pythonProject\\dataraw\\本院查明\\根据x判决如下.txt\\")
    newfileNames = ""
    with open("keywords/依照x判决如下.txt", 'r', encoding='utf-8') as fp:
        newfileNames = fp.readlines()
    l = 1
    for file in newfileNames:
        # 注意SaveAs会打开保存后的文件，有时可能看不到，但后台一定是打开的
        # path = "D:\\Code\\pythoncode\\pythonProject\\dataraw\\本院查明\\根据x判决如下.txt\\" + file.split(".docx")[0] + ".docx"
        path = "D:\\Code\\pythoncode\\pythonProject\\data_\\经审理查明\\" + file.strip("\n")
        res_path = "res.txt"
        fact = ""
        laws = []
        # print(path)
        # 获取文档的所有段落 path : 相对路径包含文档名称
        docx_temp = docx.Document(path)
        # f = False
        for para in docx_temp.paragraphs:
            if re.findall("事实(.)理由", str(para.text)):
                key1 = re.findall("事实(.)理由", str(para.text))[0]
                fact = str(para.text).split("事实" + key1 + "理由")[1]
                # if len(fact) <= 20:
                #     f = True
                #     continue
                # print(str(len(fact))+"<::>"+str(l)+"<:>"+fact)
            # if f:
            #     fact = str(para.text)
            # f = False
            # print(str(len(fact)) + "<::>" + str(l) + "<:>" + fact)
            if re.findall("依照(.*)判决如下", str(para.text)):
                t = re.findall("依照(.*)判决如下", str(para.text))[0]
                law_list = re.findall(r'[《](.*?)[》]', t)
                for i in range(len(law_list)):
                    if law_list[i] in lines:
                        laws.append(law_list[i])
        logger.info(str(l) + "<:>" + str(file))
        l += 1
        res = {'fact': fact, 'laws': laws}
        res2json = json.dumps(res, ensure_ascii=False)
        with open(res_path, 'a+', encoding='utf-8') as fp:
            fp.write(res2json + '\n')
            fp.close()


# 找到含有消费权益保护法的docx
def data_find_xf():
    newfileNames = os.listdir(r"D:\\Code\\pythoncode\\pythonProject\\data_xf\\")
    l = 1
    e = 1
    for file in newfileNames:
        try:
            # 注意SaveAs会打开保存后的文件，有时可能看不到，但后台一定是打开的
            path = "D:\\Code\\pythoncode\\pythonProject\\data_xf\\" + file.split(".docx")[0] + ".docx"
            # 获取文档的所有段落 path : 相对路径包含文档名称
            docx_temp = docx.Document(path)
            for para in docx_temp.paragraphs:
                # if re.findall("事实(.)理由", str(para.text)):
                #     key1 = re.findall("事实(.)理由", str(para.text))[0]
                #     fact = str(para.text).split("事实"+key1+"理由")[1]
                # print("["+str(l)+"]"+fact)
                # l += 1
                if re.findall("(.*)判决如下", str(para.text)):
                    t = re.findall("(.*)判决如下", str(para.text))[0]
                    law_list = re.findall(r'[《](.*?)[》]', t)
                    if "中华人民共和国消费者权益保护法" in law_list:
                        with open("keywords/中华人民共和国消费者权益保护法.txt", 'a+', encoding='utf-8') as fp:
                            fp.write(str(file) + '\n')
                            fp.close()
                        logger.info("[" + str(l) + "]" + str(file))
                        l += 1
                    break
                    # print(str(l) + "<:>" + str(re.findall("(.*)判决如下", str(para.text))))
                    # l += 1
        except:
            with open("keywords/中华人民共和国消费者权益保护法_error.txt", 'a+', encoding='utf-8') as fp:
                fp.write(str(file) + '\n')
                fp.close()
            logger.info("[" + str(e) + "]" + str(file))
            e += 1
            continue


def data_find():
    newfileNames = os.listdir(r"D:\\Code\\pythoncode\\pythonProject\\data\\")
    l = 1
    for file in newfileNames:
        path = "D:\\Code\\pythoncode\\pythonProject\\data\\" + file.split(".docx")[0] + ".docx"
        # 获取文档的所有段落 path : 相对路径包含文档名称
        docx_temp = docx.Document(path)
        f = False
        for para in docx_temp.paragraphs:
            if re.findall("查明(.*)", str(para.text)):
                key1 = re.findall("查明(.*)", str(para.text))[0]
                # fact = str(para.text).split("事实" + key1 + "理由")[1]
                f = True
            if f:
                if re.findall("(.*)判决如下", str(para.text)):
                    with open("keywords/查明.txt", 'a+', encoding='utf-8') as fp:
                        fp.write(str(file) + '\n')
                        fp.close()
                    logger.info("[" + str(l) + "]" + str(file))
                    l += 1
                    break


# 打印
def data_print():
    newfileNames = os.listdir(r"dataraw/中华人民共和国消费者权益保护法判决法条/事实x理由/依据x判决如下/")
    l = 1
    e = 1
    for file in newfileNames:
        # try:
        # 注意SaveAs会打开保存后的文件，有时可能看不到，但后台一定是打开的
        path = "dataraw/中华人民共和国消费者权益保护法判决法条/事实x理由/依据x判决如下/" + file.split(".docx")[0] + ".docx"
        # 获取文档的所有段落 path : 相对路径包含文档名称
        docx_temp = docx.Document(path)
        for para in docx_temp.paragraphs:
            # if re.findall("确认事实如下(.*)", str(para.text)):
            #     with open("keywords/确认事实如下.txt", 'a+', encoding='utf-8') as fp:
            #         fp.write(str(file) + '\n')
            #         fp.close()
            #     logger.info("[" + str(l) + "]" + str(file))
            #     l += 1
            #     break
            #     key1 = re.findall("事实(.)理由", str(para.text))[0]
            # fact = str(para.text).split("事实" + key1 + "理由")[1]
            # print("[" + str(l) + "]" +str(file))
            # l += 1
            # break
            if re.findall("(.*)判决如下", str(para.text)):
                t = re.findall("(.*)判决如下", str(para.text))[0]
                # law_list = re.findall(r'[《](.*?)[》]', t)
                #     if "中华人民共和国消费者权益保护法" in law_list:
                #     with open("keywords/中华人民共和国消费者权益保护法/综上x判决如下.txt", 'a+', encoding='utf-8') as fp:
                #         fp.write(str(file) + '\n')
                #         fp.close()
                #     logger.info("[" + str(l) + "]" + str(file))
                #     l += 1
                #     break
                #     print(str(l) + "<:>" + str(re.findall("根据(.*)判决如下", str(para.text))))
                print(str(l) + "<:>" + str(re.findall("(.*)判决如下", str(para.text))[0]))
                l += 1
    # except:
    #     with open("keywords/中华人民共和国消费者权益保护法_print_error.txt", 'a+', encoding='utf-8') as fp:
    #         fp.write(str(file) + '\n')
    #         fp.close()
    #     logger.info("[" + str(e) + "]" + str(file))
    #     e += 1
    #     continue


if __name__ == "__main__":
    # docx_transform_txt()
    # data_print()
    # data_find()
    # data_clean_sl()
    data_clean_ly()
